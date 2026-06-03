# ============================================================
# Docx 导出器 (Python)
# 通过 python-docx 精确生成符合官方模板的 Word 文档
# ============================================================

"""
Docx 导出模块。

接收 TS 端传来的中间数据，生成符合竞赛官方模板的 Word 文档。
"""

from typing import Optional, Any
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import json
import os
import base64
import io


class DocxExporter:
    """Docx 导出器"""

    def __init__(self):
        self.doc: Optional[Document] = None

    def export(self, data: dict, output_path: str) -> dict:
        """
        生成 docx 文件

        Args:
            data: 文档数据（来自 TS 端的 DocxJobData）
            output_path: 输出文件路径

        Returns:
            dict: { success: bool, output_path: str, file_size: int }
        """
        self.doc = Document()

        # 设置页面
        constraints = data.get("constraints", {})
        self._setup_page(constraints)

        # 添加标题
        template_name = data.get("templateName", "申报书")
        self._add_title(template_name)

        # 加载图片数据
        self._images = data.get("images", []) or []

        # 添加各章节
        sections = data.get("sections", [])
        for section in sections:
            self._add_section(section, constraints)

        # 保存
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        self.doc.save(output_path)

        file_size = os.path.getsize(output_path)

        return {
            "success": True,
            "output_path": output_path,
            "file_size": file_size,
        }

    def _setup_page(self, constraints: dict):
        """设置页面格式"""
        section = self.doc.sections[0]
        margins = constraints.get("margins", [2.5, 2.0, 2.5, 2.0])
        section.top_margin = Cm(margins[0])
        section.right_margin = Cm(margins[1])
        section.bottom_margin = Cm(margins[2])
        section.left_margin = Cm(margins[3])

    def _add_title(self, title: str):
        """添加文档标题"""
        heading = self.doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.size = Pt(16)
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    def _add_section(self, section: dict, constraints: dict):
        """添加章节"""
        section_id = section.get("id", "")
        section_title = section.get("title", "")
        section_type = section.get("type", "text")
        content = section.get("content", "")

        # 跳过封面/签名等特殊章节
        if section_type in ("signature",):
            return

        # 添加章节标题
        heading = self.doc.add_heading(section_title, level=1)
        for run in heading.runs:
            run.font.size = Pt(14)
            run.font.name = constraints.get("headingFont", "黑体")
            run._element.rPr.rFonts.set(qn("w:eastAsia"), constraints.get("headingFont", "黑体"))

        # 添加内容
        if section_type == "table":
            self._add_table(content, section)
        elif section_type == "abstract":
            self._add_abstract(content)
        elif section_type == "image":
            self._add_section_image(section_id, constraints)
        elif section_type == "image_group":
            self._add_section_image_group(section_id, constraints)
        else:
            self._add_text(content, constraints)

    def _add_text(self, content: Any, constraints: dict):
        """添加正文文本"""
        text = content if isinstance(content, str) else json.dumps(content, ensure_ascii=False)

        # 按段落拆分
        paragraphs = text.split("\n")
        for para_text in paragraphs:
            if para_text.strip():
                p = self.doc.add_paragraph()
                run = p.add_run(para_text.strip())
                run.font.size = Pt(12)
                run.font.name = constraints.get("font", "宋体")
                run._element.rPr.rFonts.set(qn("w:eastAsia"), constraints.get("font", "宋体"))
                p.paragraph_format.line_spacing = constraints.get("lineSpacing", 1.5)

    def _add_abstract(self, content: Any):
        """添加摘要（斜体）"""
        text = content if isinstance(content, str) else json.dumps(content, ensure_ascii=False)
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(11)

    def _add_section_image(self, section_id: str, constraints: dict):
        """插入单张图片"""
        images_for_section = [img for img in self._images if img.get("sectionId") == section_id]
        for img in images_for_section:
            self._insert_image(img, constraints)
            if img.get("label"):
                self._add_image_caption(img["label"], constraints)

    def _add_section_image_group(self, section_id: str, constraints: dict):
        """插入多张图片（并排）"""
        images_for_section = [img for img in self._images if img.get("sectionId") == section_id]
        if not images_for_section:
            return

        # 多图并排（每行最多2张）
        for i in range(0, len(images_for_section), 2):
            row_images = images_for_section[i:i+2]
            self._add_image_row(row_images, constraints)

    def _insert_image(self, img: dict, constraints: dict):
        """插入单张图片到文档"""
        try:
            # 支持 base64 和文件路径
            image_data = img.get("imageData")
            if image_data:
                # base64 解码
                raw = base64.b64decode(image_data)
                image_stream = io.BytesIO(raw)
            else:
                file_path = img.get("filePath", "")
                if not os.path.exists(file_path):
                    return
                image_stream = file_path

            # 计算尺寸
            max_width_cm = self._parse_cm(img.get("maxWidth", "14"))
            max_height_cm = self._parse_cm(img.get("maxHeight", "20"))

            paragraph = self.doc.add_paragraph()
            paragraph.alignment = {
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }.get(img.get("position", "center"), WD_ALIGN_PARAGRAPH.CENTER)

            run = paragraph.add_run()
            if isinstance(image_stream, io.BytesIO):
                run.add_picture(image_stream, width=Cm(max_width_cm))
            else:
                run.add_picture(image_stream, width=Cm(max_width_cm))

        except Exception as e:
            # 图片插入失败时，添加占位文本
            p = self.doc.add_paragraph()
            p.add_run(f"[图片: {img.get('label', '未命名')}]").italic = True

    def _add_image_row(self, row_images: list, constraints: dict):
        """并排插入多张图片"""
        if not row_images:
            return

        # 创建表格（1行N列）来并排放置图片
        table = self.doc.add_table(rows=1, cols=len(row_images))
        table.style = "Table Grid"

        for i, img in enumerate(row_images):
            cell = table.rows[0].cells[i]
            # 清空默认段落
            cell.text = ""
            try:
                image_data = img.get("imageData")
                if image_data:
                    raw = base64.b64decode(image_data)
                    image_stream = io.BytesIO(raw)
                else:
                    file_path = img.get("filePath", "")
                    if not os.path.exists(file_path):
                        continue
                    image_stream = file_path

                width_cm = self._parse_cm(img.get("maxWidth", "8"))
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                if isinstance(image_stream, io.BytesIO):
                    run.add_picture(image_stream, width=Cm(width_cm))
                else:
                    run.add_picture(image_stream, width=Cm(width_cm))

                if img.get("label"):
                    caption_p = cell.add_paragraph()
                    caption_p.add_run(img["label"]).font.size = Pt(9)

            except Exception:
                cell.text = f"[{img.get('label', '图')}]"

    def _add_image_caption(self, text: str, constraints: dict):
        """添加图注"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = constraints.get("font", "宋体")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), constraints.get("font", "宋体"))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)

    def _parse_cm(self, value: Any) -> float:
        """解析厘米值，支持 '14cm' 或纯数字"""
        if isinstance(value, (int, float)):
            return float(value)
        s = str(value).replace("cm", "").strip()
        return float(s)

    def _add_table(self, content: Any, section: dict):
        """添加表格"""
        if not isinstance(content, list) or len(content) == 0:
            return

        # 获取表头
        headers = list(content[0].keys()) if isinstance(content[0], dict) else []

        table = self.doc.add_table(rows=len(content) + 1, cols=max(len(headers), 1))
        table.style = "Table Grid"

        # 表头
        if headers:
            for j, header in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        # 数据行
        for i, row_data in enumerate(content):
            if isinstance(row_data, dict):
                for j, key in enumerate(headers):
                    if j < len(table.rows[i + 1].cells):
                        table.rows[i + 1].cells[j].text = str(row_data.get(key, ""))


def export_docx(data: dict, output_path: str) -> dict:
    """便捷函数"""
    exporter = DocxExporter()
    return exporter.export(data, output_path)


class DocxExporterAgent:
    """Docx 导出器 Agent（注册到桥服务）"""

    name = "docx-exporter"
    description = "使用 python-docx 导出 Word 文档"

    def export(self, data: dict, output_path: str) -> dict:
        """导出 docx"""
        exporter = DocxExporter()
        return exporter.export(data, output_path)

    def run(self, **params) -> dict:
        """通用的 run 方法"""
        data = params.get("data", {})
        output_path = params.get("output_path", "output.docx")
        return self.export(data, output_path)

    def ping(self) -> dict:
        return {"pong": True}
