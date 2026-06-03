from __future__ import annotations

import json
import os
import re
import sys
import importlib
import tempfile
import base64
import urllib.request
import urllib.error
from pathlib import Path


def split_markdown(markdown: str):
    blocks = []
    table_lines = []
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            continue
        if table_lines:
            blocks.append(("table", table_lines))
            table_lines = []
        if not line.strip():
            blocks.append(("blank", ""))
        elif line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif re.match(r"^!\[(.*?)\]\((.*?)\)$", line):
            match = re.match(r"^!\[(.*?)\]\((.*?)\)$", line)
            blocks.append(("image", {"alt": match.group(1).strip(), "src": match.group(2).strip()}))
        elif re.match(r"^[-*]\s+", line):
            blocks.append(("bullet", re.sub(r"^[-*]\s+", "", line).strip()))
        elif re.match(r"^\d+\.\s+", line):
            blocks.append(("number", line.strip()))
        else:
            blocks.append(("p", line.strip()))
    if table_lines:
        blocks.append(("table", table_lines))
    return blocks


def parse_table(lines: list[str]):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            continue
        rows.append(cells)
    width = max((len(row) for row in rows), default=0)
    return [row + [""] * (width - len(row)) for row in rows]


def clean_inline(text: str):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"([A-Za-z])\s+-\s+([A-Za-z])", r"\1-\2", text)
    text = re.sub(r"([A-Za-z])\s+/\s+([A-Za-z])", r"\1/\2", text)
    text = re.sub(r"\b(Phase)\s+([AB])\s*/\s*(Phase)\s+([AB])\b", r"\1 \2/\3 \4", text)
    text = re.sub(r"\b(HMAD)\s*-\s*(Ednet)\b", r"\1-\2", text)
    text = re.sub(r"\b(SPA)\s*-\s*(HyperNet)\b", r"\1-\2", text)
    text = re.sub(r"\b(Reptile)\s+元", r"\1元", text)
    return text


def cell_display_len(text: str):
    value = clean_inline(str(text or "")).replace("\n", "")
    return sum(2 if "\u4e00" <= char <= "\u9fff" else 1 for char in value)


def soft_wrap_cell_text(text: str, limit: int = 22):
    value = clean_inline(str(text or "")).replace("\r", "").strip()
    if not value or "\n" in value or cell_display_len(value) <= limit:
        return value
    lines: list[str] = []
    current = ""
    current_len = 0
    for char in value:
        char_len = 2 if "\u4e00" <= char <= "\u9fff" else 1
        current += char
        current_len += char_len
        if current_len >= limit and char in "，,；;。！？、 ":
            lines.append(current.strip())
            current = ""
            current_len = 0
        elif current_len >= limit + 4:
            lines.append(current.strip())
            current = ""
            current_len = 0
    if current.strip():
        lines.append(current.strip())
    return "\n".join(lines)


def pdf_paragraph_cell_text(text: str, limit: int):
    from xml.sax.saxutils import escape

    return escape(soft_wrap_cell_text(text, limit)).replace("\n", "<br/>")


def table_col_weights(rows: list[list[str]]):
    if not rows:
        return []
    cols = max((len(row) for row in rows), default=0)
    weights: list[float] = []
    for col in range(cols):
        values = [row[col] if col < len(row) else "" for row in rows]
        header_len = cell_display_len(values[0] if values else "")
        max_len = max((cell_display_len(value) for value in values[:16]), default=6)
        weights.append(max(7.0, min(22.0, max(header_len * 1.25, max_len * 0.72, 8.0))))
    return weights


def table_col_widths_dxa(rows: list[list[str]], total_dxa: int = 8500):
    weights = table_col_weights(rows)
    if not weights:
        return []
    total_weight = sum(weights) or len(weights)
    return [int(total_dxa * weight / total_weight) for weight in weights]


def table_col_widths_pdf(rows: list[list[str]], total_width: float):
    weights = table_col_weights(rows)
    if not weights:
        return None
    total_weight = sum(weights) or len(weights)
    return [total_width * weight / total_weight for weight in weights]


def strip_submission_only_markdown(markdown: str):
    """Keep only the formal project-book body before exporting."""
    text = str(markdown or "").strip()
    submission_heading = re.compile(
        r"(?m)^#{1,6}\s*("
        r"项目书评审返修报告|终稿质量检测报告|终稿质量检测|质量体检报告|"
        r"材料来源与正文对应表|自动去重修稿说明|评审返修落实说明|"
        r"用户自定义产物要求核对|附录A[^\n]*|附录B[^\n]*|附录C[^\n]*"
        r")\s*$"
    )
    cut_points = [match.start() for match in submission_heading.finditer(text)]
    if cut_points:
        text = text[: min(cut_points)].rstrip()
    blocked = re.compile(
        r"(质量报告|质检报告|质量体检|评审返修报告|系统说明|来源映射|Paper-agent\s*负责|"
        r"自动修稿|自动生成|当前章节|写作建议|修改建议|待补充|后续完善|"
        r"以实际提交附件为准|可考虑|如有条件|TODO|\?\?\?)"
    )
    kept: list[str] = []
    for line in text.splitlines():
        if blocked.search(line):
            continue
        kept.append(line)
    return re.sub(r"\n{3,}", "\n\n", "\n".join(kept)).strip()


def text_value(project: dict, key: str, fallback: str):
    value = str(project.get(key) or "").strip()
    return value if value else fallback


def template_profile(project: dict):
    template = str(project.get("template") or project.get("competition") or "").lower()
    doc_style = str(project.get("docStyle") or "competition").lower()
    if doc_style == "nature":
        return {
            "id": "nature",
            "label": "Nature 学术风格",
            "body_font": "宋体",
            "heading_font": "黑体",
            "body_size": 11,
            "body_line": 1.35,
            "pdf_leading": 18,
            "h1_size": 16,
            "h2_size": 14,
            "h3_size": 12.5,
            "caption_size": 9.5,
            "table_size": 8.8,
            "table_leading": 11,
            "margins": {"top": 2.3, "bottom": 2.3, "left": 2.6, "right": 2.2},
            "header": "学术项目书",
            "table_fill": "F3F4F6",
            "line_color": "9CA3AF",
        }
    if "internet" in template or "互联网" in template or "plus" in template:
        return {
            "id": "internet-plus",
            "label": "互联网+ 商业计划书",
            "body_font": "宋体",
            "heading_font": "黑体",
            "body_size": 12,
            "body_line": 1.45,
            "pdf_leading": 21,
            "h1_size": 16,
            "h2_size": 15,
            "h3_size": 14,
            "caption_size": 10.5,
            "table_size": 9.5,
            "table_leading": 13,
            "margins": {"top": 2.4, "bottom": 2.4, "left": 2.8, "right": 2.3},
            "header": "互联网+ 商业计划书",
            "table_fill": "EEF2FF",
            "line_color": "A5B4FC",
        }
    if "tiaozhan" in template or "challenge" in template or "挑战" in template:
        return {
            "id": "tiaozhanbei",
            "label": "挑战杯项目计划书",
            "body_font": "宋体",
            "heading_font": "黑体",
            "body_size": 12,
            "body_line": 1.5,
            "pdf_leading": 22,
            "h1_size": 16,
            "h2_size": 15,
            "h3_size": 14,
            "caption_size": 10.5,
            "table_size": 9.5,
            "table_leading": 13,
            "margins": {"top": 2.54, "bottom": 2.54, "left": 3.0, "right": 2.5},
            "header": "挑战杯项目计划书",
            "table_fill": "EAF1F8",
            "line_color": "94A3B8",
        }
    return {
        "id": "dachuang",
        "label": "大创项目计划书",
        "body_font": "宋体",
        "heading_font": "黑体",
        "body_size": 12,
        "body_line": 1.5,
        "pdf_leading": 22,
        "h1_size": 16,
        "h2_size": 15,
        "h3_size": 14,
        "caption_size": 10.5,
        "table_size": 9.5,
        "table_leading": 13,
        "margins": {"top": 2.54, "bottom": 2.54, "left": 3.0, "right": 2.5},
        "header": "大学生创新创业训练计划项目书",
        "table_fill": "EAF1F8",
        "line_color": "CBD5E1",
    }


def load_cjk_font(size: int, bold: bool = False):
    from PIL import ImageFont

    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf" if bold else r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arialuni.ttf",
    ]
    for font_path in candidates:
        if font_path and os.path.exists(font_path):
            try:
                return ImageFont.truetype(font_path, size)
            except Exception:
                pass
    return ImageFont.load_default()


def ellipsize(text: str, limit: int):
    text = re.sub(r"\s+", " ", text).strip()
    return text if len(text) <= limit else text[: limit - 1] + "…"


def wrap_text(draw, text: str, font, max_width: int, max_lines: int = 3):
    chars = list(text)
    lines: list[str] = []
    current = ""
    for char in chars:
        trial = current + char
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
            continue
        if current:
            lines.append(current)
        current = char
        if len(lines) >= max_lines:
            break
    if current and len(lines) < max_lines:
        lines.append(current)
    if len(lines) == max_lines and len("".join(lines)) < len(text):
        lines[-1] = lines[-1].rstrip("，。；、 ") + "…"
    return lines


def draw_centered_text(draw, box, text: str, font, fill="#111827", max_lines: int = 3):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 24, max_lines=max_lines)
    line_height = font.size + 6
    total = len(lines) * line_height
    y = y1 + ((y2 - y1) - total) / 2
    for line in lines:
        width = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=font, fill=fill)
        y += line_height


def arrow(draw, start, end, fill="#2563eb", width=4):
    import math

    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 12
    points = [
        end,
        (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6)),
        (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6)),
    ]
    draw.polygon(points, fill=fill)


def diagram_path(kind: str, alt: str, project: dict, output_path: str):
    from PIL import Image, ImageDraw

    temp_dir = Path(output_path).parent / "_paper_figures"
    temp_dir.mkdir(parents=True, exist_ok=True)
    target = temp_dir / f"{kind}.png"

    width, height = 1400, 820
    image = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(image)
    title_font = load_cjk_font(42, bold=True)
    label_font = load_cjk_font(28, bold=True)
    small_font = load_cjk_font(22)
    note_font = load_cjk_font(20)

    project_name = ellipsize(text_value(project, "name", "项目"), 28)
    track = ellipsize(text_value(project, "track", "目标应用场景"), 24)
    product = ellipsize(text_value(project, "product", "核心算法/系统原型/服务平台"), 42)
    market = ellipsize(text_value(project, "market", "目标客户与试点场景"), 36)

    draw.text((70, 48), alt, font=title_font, fill="#1d4ed8")
    draw.text((70, 103), project_name, font=note_font, fill="#64748b")

    if kind == "architecture":
        boxes = [
            ((80, 220, 300, 370), "场景/数据输入", track, "#eff6ff"),
            ((365, 220, 585, 370), "场景路由模块", "识别任务类型与场景差异", "#f0fdf4"),
            ((650, 220, 870, 370), "核心检测网络", product, "#fff7ed"),
            ((935, 220, 1155, 370), "应用服务层", "预警、报告、接口与权限", "#f8fafc"),
            ((1080, 500, 1300, 650), "展示与交付", market, "#eef2ff"),
            ((650, 500, 870, 650), "元适应优化", "反馈样本、指标复盘、持续迭代", "#fdf2f8"),
        ]
        for box, title, body, color in boxes:
            draw.rounded_rectangle(box, radius=22, fill=color, outline="#cbd5e1", width=3)
            draw_centered_text(draw, (box[0], box[1] + 12, box[2], box[1] + 64), title, label_font, "#0f172a", 1)
            draw_centered_text(draw, (box[0] + 12, box[1] + 70, box[2] - 12, box[3] - 12), body, small_font, "#334155", 3)
        arrow(draw, (300, 295), (365, 295))
        arrow(draw, (585, 295), (650, 295))
        arrow(draw, (870, 295), (935, 295))
        arrow(draw, (1045, 370), (1135, 500))
        arrow(draw, (1080, 575), (870, 575))
        arrow(draw, (760, 500), (760, 370))
        draw.text((80, 720), "说明：技术架构围绕项目名称、赛道、产品模块和目标市场组织，可按真实模块名称更新。", font=note_font, fill="#64748b")
    else:
        steps = [
            ("需求调研", "访谈/问卷/场景确认"),
            ("方案设计", "技术路线与交付边界"),
            ("原型验证", "核心功能和指标测试"),
            ("试点部署", "真实或模拟场景运行"),
            ("验收交付", "报告、培训、记录沉淀"),
            ("运维迭代", "反馈复盘与版本升级"),
        ]
        y = 250
        gap = 42
        box_w, box_h = 180, 142
        for index, (title, body) in enumerate(steps):
            x = 70 + index * (box_w + gap)
            box = (x, y, x + box_w, y + box_h)
            color = "#eff6ff" if index % 2 == 0 else "#f0fdf4"
            draw.rounded_rectangle(box, radius=20, fill=color, outline="#cbd5e1", width=3)
            draw_centered_text(draw, (box[0], box[1] + 14, box[2], box[1] + 58), title, label_font, "#0f172a", 1)
            draw_centered_text(draw, (box[0] + 12, box[1] + 66, box[2] - 12, box[3] - 12), body, small_font, "#334155", 2)
            if index < len(steps) - 1:
                arrow(draw, (box[2], y + box_h // 2), (box[2] + gap - 8, y + box_h // 2))
        draw.rounded_rectangle((210, 520, 1190, 650), radius=24, fill="#f8fafc", outline="#cbd5e1", width=3)
        draw_centered_text(draw, (230, 535, 1170, 635), f"闭环目标：围绕{track}形成可验证、可交付、可复盘的项目服务流程", small_font, "#334155", 2)
        draw.text((80, 720), "说明：该流程图适用于挑战杯/大创项目书中的服务实施计划、商业落地路径和交付闭环说明。", font=note_font, fill="#64748b")

    image.save(target)
    return str(target)


def external_image_prompt(kind: str, alt: str, project: dict):
    project_name = text_value(project, "name", "项目")
    track = text_value(project, "track", "应用场景")
    product = text_value(project, "product", "核心产品")
    market = text_value(project, "market", "目标客户")
    if kind == "architecture":
        return (
            f"为挑战杯/大创项目书生成一张清晰的中文技术架构图。标题：{alt}。"
            f"项目名称：{project_name}。赛道/场景：{track}。核心产品：{product}。目标市场：{market}。"
            "风格要求：白底、蓝灰专业配色、流程框图、模块清晰、适合放入项目计划书，包含数据输入、核心算法/系统、应用服务、反馈优化、交付展示等模块。"
        )
    return (
        f"为挑战杯/大创项目书生成一张中文服务实施流程图。标题：{alt}。"
        f"项目名称：{project_name}。赛道/场景：{track}。目标市场：{market}。"
        "风格要求：白底、蓝灰专业配色、横向流程图、步骤清晰，包含需求调研、方案设计、原型验证、试点部署、验收交付、运维迭代。"
    )


def external_image_path(kind: str, alt: str, project: dict, output_path: str):
    api_key = str(project.get("imageApiKey") or "").strip()
    base_url = str(project.get("imageBaseUrl") or "").strip().rstrip("/")
    model = str(project.get("imageModel") or "gpt-image-1").strip()
    if not api_key or not base_url or not model:
        return ""

    temp_dir = Path(output_path).parent / "_paper_figures"
    temp_dir.mkdir(parents=True, exist_ok=True)
    target = temp_dir / f"{kind}-external.png"
    payload = json.dumps({
        "model": model,
        "prompt": external_image_prompt(kind, alt, project),
        "size": "1024x1024",
        "n": 1,
    }).encode("utf-8")
    request = urllib.request.Request(
        f"{base_url}/v1/images/generations",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=45) as response:
            data = json.loads(response.read().decode("utf-8"))
        item = (data.get("data") or [{}])[0]
        if item.get("b64_json"):
            target.write_bytes(base64.b64decode(item["b64_json"]))
            return str(target)
        if item.get("url"):
            with urllib.request.urlopen(item["url"], timeout=45) as image_response:
                target.write_bytes(image_response.read())
            return str(target)
    except Exception:
        return ""
    return ""


def resolve_image(value: dict, project: dict, output_path: str):
    src = value.get("src", "")
    alt = value.get("alt", "图示")
    if src == "paper://figure/architecture":
        external = external_image_path("architecture", alt, project, output_path)
        if external:
            return external
        return diagram_path("architecture", alt, project, output_path)
    if src == "paper://figure/service-flow":
        external = external_image_path("service-flow", alt, project, output_path)
        if external:
            return external
        return diagram_path("service-flow", alt, project, output_path)
    if os.path.exists(src):
        return src
    return ""


def import_python_docx():
    # The legacy exporter in this folder is named docx.py, which would shadow
    # the third-party python-docx package when this script is executed directly.
    local_dir = str(Path(__file__).resolve().parent)
    original_path = list(sys.path)
    sys.path = [entry for entry in sys.path if str(Path(entry or ".").resolve()) != local_dir]
    cached = sys.modules.get("docx")
    if cached is not None:
        cached_file = getattr(cached, "__file__", "")
        if cached_file and str(Path(cached_file).resolve()).startswith(local_dir):
            del sys.modules["docx"]
    try:
        return importlib.import_module("docx")
    finally:
        sys.path = original_path


def set_docx_run_font(run, font_name: str, size_pt: float, bold: bool = False):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_docx_paragraph_format(paragraph, *, first_line: bool = True, align=None, line_spacing: float = 1.5, before: float = 0, after: float = 0):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    # 小四正文约 12pt，首行缩进 2 个中文字符即约 24pt。
    paragraph.paragraph_format.first_line_indent = Pt(24) if first_line else Pt(0)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def add_docx_page_number(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph.alignment = 1
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_text, fld_end])
    set_docx_run_font(run, "宋体", 10.5)


def set_docx_cell_shading(cell, fill: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_docx_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_docx_table_width(table, width_dxa: int = 8500):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "autofit")


def set_docx_cell_width(cell, width_dxa: int):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_docx_horizontal_line(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D1D5DB")
    borders.append(bottom)


def configure_docx_styles(doc, project: dict):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    profile = template_profile(project)
    margins = profile["margins"]
    section = doc.sections[0]
    section.top_margin = Cm(margins["top"])
    section.bottom_margin = Cm(margins["bottom"])
    section.left_margin = Cm(margins["left"])
    section.right_margin = Cm(margins["right"])
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), profile["body_font"])
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(profile["body_size"])
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = profile["body_line"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    heading_specs = {
        "Heading 1": (profile["heading_font"], profile["h1_size"], True, WD_ALIGN_PARAGRAPH.CENTER, 12, 8, 0),
        "Heading 2": (profile["heading_font"], profile["h2_size"], True, WD_ALIGN_PARAGRAPH.LEFT, 10, 6, 1),
        "Heading 3": (profile["heading_font"], profile["h3_size"], True, WD_ALIGN_PARAGRAPH.LEFT, 8, 4, 2),
    }
    for style_name, (font_name, size, bold, align, before, after, outline_level) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.alignment = align
        style.paragraph_format.line_spacing = 1.3
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        p_pr = style._element.get_or_add_pPr()
        outline = p_pr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            p_pr.append(outline)
        outline.set(qn("w:val"), str(outline_level))

    header = section.header.paragraphs[0]
    header.text = f"{text_value(project, 'name', '项目计划书')}｜{profile['header']}"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        set_docx_run_font(run, profile["body_font"], 10.5)
    add_docx_horizontal_line(header)

    footer = section.footer.paragraphs[0]
    add_docx_page_number(footer)


def export_docx(markdown: str, output_path: str, project: dict):
    docx_module = import_python_docx()
    Document = docx_module.Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Pt, Inches

    doc = Document()
    configure_docx_styles(doc, project)
    profile = template_profile(project)
    figure_index = 1
    table_index = 1
    last_heading = ""
    toc_open = False

    def add_page_break():
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    for kind, value in split_markdown(markdown):
        if kind == "blank":
            continue
        if kind == "image":
            image_path = resolve_image(value, project, output_path)
            if image_path:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(6.2))
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption.add_run(f"图{figure_index} {clean_inline(value.get('alt', '项目图示'))}")
                set_docx_run_font(caption_run, profile["body_font"], profile["caption_size"])
                caption.paragraph_format.space_after = Pt(6)
                figure_index += 1
            continue
        if kind in {"h1", "h2", "h3"}:
            heading_text = clean_inline(value)
            if toc_open and heading_text != "目录":
                add_page_break()
                toc_open = False
            level = {"h1": 0, "h2": 1, "h3": 2}[kind]
            paragraph = doc.add_heading(heading_text, level=level)
            paragraph.style = {"h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3"}[kind]
            for run in paragraph.runs:
                size = profile["h1_size"] if kind == "h1" else profile["h2_size"] if kind == "h2" else profile["h3_size"]
                set_docx_run_font(run, profile["heading_font"], size, bold=True)
            last_heading = heading_text
            if heading_text == "目录":
                toc_open = True
            continue
        if kind == "table":
            rows = parse_table(value)
            if not rows:
                continue
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run(f"表{table_index} {clean_inline(rows[0][0]) if rows and rows[0] else '项目表格'}")
            set_docx_run_font(caption_run, profile["body_font"], profile["caption_size"], bold=True)
            caption.paragraph_format.space_before = Pt(6)
            caption.paragraph_format.space_after = Pt(4)
            table_index += 1

            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            total_width = 8500
            col_widths = table_col_widths_dxa(rows, total_width)
            set_docx_table_width(table, total_width)
            for i, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    cell = table.cell(i, j)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    if j < len(col_widths):
                        set_docx_cell_width(cell, col_widths[j])
                    set_docx_cell_margins(cell)
                    if i == 0:
                        set_docx_cell_shading(cell, profile["table_fill"])
                    wrap_limit = 16 if len(rows[0]) >= 5 else 22 if len(rows[0]) >= 4 else 28
                    cell.text = soft_wrap_cell_text(cell_text, wrap_limit)
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                        paragraph.paragraph_format.first_line_indent = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.15
                        paragraph.paragraph_format.space_after = Pt(0)
                        for run in paragraph.runs:
                            set_docx_run_font(run, profile["body_font"], profile["table_size"], bold=i == 0)
            if last_heading == "封面信息":
                add_page_break()
            continue
        paragraph = doc.add_paragraph()
        if kind == "bullet":
            run = paragraph.add_run(clean_inline(value))
            set_docx_paragraph_format(paragraph, first_line=False, line_spacing=profile["body_line"])
            paragraph.paragraph_format.left_indent = Pt(24)
            paragraph.paragraph_format.first_line_indent = Pt(0)
        elif kind == "number":
            run = paragraph.add_run(clean_inline(value))
            set_docx_paragraph_format(paragraph, first_line=False, line_spacing=profile["body_line"])
            paragraph.paragraph_format.left_indent = Pt(24)
            paragraph.paragraph_format.first_line_indent = Pt(0)
        else:
            run = paragraph.add_run(clean_inline(value))
            set_docx_paragraph_format(paragraph, first_line=True, line_spacing=profile["body_line"])
        set_docx_run_font(run, profile["body_font"], profile["body_size"])

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def register_pdf_font():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\arialuni.ttf",
    ]
    for font_path in candidates:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont("CJK", font_path))
                return "CJK"
            except Exception:
                pass
    return "Helvetica"


def export_pdf(markdown: str, output_path: str, project: dict):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Image as PdfImage, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font = register_pdf_font()
    profile = template_profile(project)
    margins = profile["margins"]
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("CNNormal", parent=styles["Normal"], fontName=font, fontSize=profile["body_size"], leading=profile["pdf_leading"], firstLineIndent=24, alignment=0, spaceAfter=0, wordWrap="CJK", splitLongWords=0)
    h1 = ParagraphStyle("CNH1", parent=styles["Title"], fontName=font, fontSize=profile["h1_size"], leading=24, alignment=1, spaceBefore=12, spaceAfter=8)
    h2 = ParagraphStyle("CNH2", parent=styles["Heading1"], fontName=font, fontSize=profile["h2_size"], leading=22, textColor=colors.black, spaceBefore=10, spaceAfter=6)
    h3 = ParagraphStyle("CNH3", parent=styles["Heading2"], fontName=font, fontSize=profile["h3_size"], leading=20, spaceBefore=8, spaceAfter=4)
    bullet = ParagraphStyle("CNBullet", parent=normal, firstLineIndent=24)
    caption = ParagraphStyle("CNCaption", parent=styles["Normal"], fontName=font, fontSize=profile["caption_size"], leading=15, alignment=1, spaceAfter=6)
    table_body = ParagraphStyle("CNTableBody", parent=styles["Normal"], fontName=font, fontSize=profile["table_size"], leading=profile["table_leading"], alignment=0, wordWrap="CJK", spaceBefore=0, spaceAfter=0)
    table_head = ParagraphStyle("CNTableHead", parent=table_body, alignment=1)

    story = []
    figure_index = 1
    table_index = 1
    last_heading = ""
    toc_open = False
    for kind, value in split_markdown(markdown):
        if kind == "blank":
            story.append(Spacer(1, 5))
        elif kind == "h1":
            heading_text = clean_inline(value)
            if toc_open and heading_text != "目录":
                story.append(PageBreak())
                toc_open = False
            story.append(Paragraph(heading_text, h1))
            last_heading = heading_text
        elif kind == "h2":
            heading_text = clean_inline(value)
            if toc_open and heading_text != "目录":
                story.append(PageBreak())
                toc_open = False
            story.append(Paragraph(heading_text, h2))
            last_heading = heading_text
            if heading_text == "目录":
                toc_open = True
        elif kind == "h3":
            heading_text = clean_inline(value)
            if toc_open and heading_text != "目录":
                story.append(PageBreak())
                toc_open = False
            story.append(Paragraph(heading_text, h3))
            last_heading = heading_text
        elif kind == "image":
            image_path = resolve_image(value, project, output_path)
            if image_path:
                story.append(PdfImage(image_path, width=16.5 * cm, height=9.67 * cm))
                story.append(Paragraph(f"图{figure_index} {clean_inline(value.get('alt', '项目图示'))}", caption))
                story.append(Spacer(1, 10))
                figure_index += 1
        elif kind == "bullet":
            story.append(Paragraph(clean_inline(value), bullet))
        elif kind == "number":
            story.append(Paragraph(clean_inline(value), normal))
        elif kind == "table":
            rows = parse_table(value)
            if rows:
                title = clean_inline(rows[0][0]) if rows and rows[0] else "项目表格"
                story.append(Paragraph(f"表{table_index} {title}", caption))
                table_index += 1
                available_width = A4[0] - (margins["left"] + margins["right"]) * cm
                col_widths = table_col_widths_pdf(rows, available_width)
                wrap_limit = 14 if len(rows[0]) >= 5 else 20 if len(rows[0]) >= 4 else 28
                table_data = [
                    [
                        Paragraph(pdf_paragraph_cell_text(cell, wrap_limit), table_head if row_index == 0 else table_body)
                        for cell in row
                    ]
                    for row_index, row in enumerate(rows)
                ]
                table = Table(table_data, colWidths=col_widths, repeatRows=1, hAlign="CENTER")
                table.setStyle(TableStyle([
                    ("FONTNAME", (0, 0), (-1, -1), font),
                    ("FONTSIZE", (0, 0), (-1, -1), profile["table_size"]),
                    ("LEADING", (0, 0), (-1, -1), profile["table_leading"]),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{profile['table_fill']}")),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{profile['line_color']}")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ]))
                story.append(table)
                story.append(Spacer(1, 8))
                if last_heading == "封面信息":
                    story.append(PageBreak())
        else:
            story.append(Paragraph(clean_inline(value), normal))

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=margins["right"] * cm,
        leftMargin=margins["left"] * cm,
        topMargin=margins["top"] * cm,
        bottomMargin=margins["bottom"] * cm,
    )

    def draw_header_footer(canvas, document):
        canvas.saveState()
        title = f"{text_value(project, 'name', '项目计划书')}｜{profile['header']}"
        canvas.setFont(font, 9)
        canvas.setFillColor(colors.HexColor("#4b5563"))
        canvas.drawCentredString(A4[0] / 2, A4[1] - 1.55 * cm, title[:48])
        canvas.setStrokeColor(colors.HexColor(f"#{profile['line_color']}"))
        canvas.setLineWidth(0.4)
        canvas.line(margins["left"] * cm, A4[1] - 1.72 * cm, A4[0] - margins["right"] * cm, A4[1] - 1.72 * cm)
        canvas.drawCentredString(A4[0] / 2, 1.45 * cm, str(document.page))
        canvas.restoreState()

    doc.build(story, onFirstPage=draw_header_footer, onLaterPages=draw_header_footer)


LATEX_SPECIALS = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def latex_escape(text: str):
    return "".join(LATEX_SPECIALS.get(char, char) for char in clean_inline(text))


def latex_image_path(path: str):
    return Path(path).resolve().as_posix().replace("%", r"\%")


def export_latex(markdown: str, output_path: str, project: dict):
    title = latex_escape(text_value(project, "name", "项目计划书"))
    lines = [
        r"\documentclass[UTF8,a4paper,12pt]{ctexart}",
        r"\usepackage{geometry}",
        r"\usepackage{graphicx}",
        r"\usepackage{booktabs}",
        r"\usepackage{longtable}",
        r"\usepackage{array}",
        r"\usepackage{setspace}",
        r"\usepackage{fancyhdr}",
        r"\usepackage{hyperref}",
        r"\geometry{left=3.0cm,right=2.5cm,top=2.54cm,bottom=2.54cm}",
        r"\setlength{\parindent}{2em}",
        r"\setlength{\parskip}{0em}",
        r"\onehalfspacing",
        r"\pagestyle{fancy}",
        r"\fancyhf{}",
        rf"\fancyhead[C]{{{title}}}",
        r"\fancyfoot[C]{\thepage}",
        rf"\title{{{title}}}",
        r"\author{}",
        r"\date{}",
        r"\begin{document}",
        r"\maketitle",
        r"\tableofcontents",
        r"\newpage",
    ]

    for kind, value in split_markdown(markdown):
        if kind == "blank":
            lines.append("")
        elif kind == "h1":
            lines.append(rf"\section{{{latex_escape(value)}}}")
        elif kind == "h2":
            lines.append(rf"\subsection{{{latex_escape(value)}}}")
        elif kind == "h3":
            lines.append(rf"\subsubsection{{{latex_escape(value)}}}")
        elif kind == "image":
            try:
                image_path = resolve_image(value, project, output_path)
            except Exception:
                image_path = ""
            if image_path:
                alt = latex_escape(value.get("alt", "图示"))
                lines.extend([
                    r"\begin{figure}[htbp]",
                    r"\centering",
                    rf"\includegraphics[width=0.92\textwidth]{{{latex_image_path(image_path)}}}",
                    rf"\caption{{{alt}}}",
                    r"\end{figure}",
                ])
            else:
                lines.append(rf"\textbf{{图示说明：{latex_escape(value.get('alt', '项目图示'))}}}")
        elif kind == "bullet":
            lines.extend([r"\begin{itemize}", rf"\item {latex_escape(value)}", r"\end{itemize}"])
        elif kind == "number":
            item = re.sub(r"^\d+\.\s+", "", value).strip()
            lines.extend([r"\begin{enumerate}", rf"\item {latex_escape(item)}", r"\end{enumerate}"])
        elif kind == "table":
            rows = parse_table(value)
            if rows:
                col_count = len(rows[0])
                col_spec = "|".join(["p{0.22\\textwidth}"] * col_count)
                lines.append(rf"\begin{{longtable}}{{{col_spec}}}")
                for index, row in enumerate(rows):
                    cells = [latex_escape(cell) for cell in row]
                    lines.append(" & ".join(cells) + r" \\")
                    lines.append(r"\hline" if index == 0 else "")
                lines.append(r"\end{longtable}")
        else:
            lines.append(latex_escape(value))

    lines.append(r"\end{document}")
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    Path(output_path).write_text("\n".join(lines), encoding="utf-8")


def main():
    payload = json.load(sys.stdin)
    markdown = strip_submission_only_markdown(payload["markdown"])
    output_path = payload["output_path"]
    fmt = payload["format"]
    project = payload.get("project") or {}
    if fmt == "docx":
        export_docx(markdown, output_path, project)
    elif fmt == "pdf":
        export_pdf(markdown, output_path, project)
    elif fmt in {"tex", "latex"}:
        export_latex(markdown, output_path, project)
    else:
        raise ValueError(f"Unsupported format: {fmt}")
    print(json.dumps({
        "success": True,
        "output_path": output_path,
        "file_size": os.path.getsize(output_path),
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
